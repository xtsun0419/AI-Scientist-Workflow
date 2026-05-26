from __future__ import annotations

import json
import os
import re
import ssl
import statistics
import subprocess
import tempfile
import urllib.error
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
SKILL_ROOT = ROOT / "academic-research-skills-main" / "academic-paper-reviewer"
AGENT_ROOT = SKILL_ROOT / "agents"


AGENT_FILES = {
    "field_analyst": "field_analyst_agent.md",
    "eic": "eic_agent.md",
    "methodology": "methodology_reviewer_agent.md",
    "domain": "domain_reviewer_agent.md",
    "perspective": "perspective_reviewer_agent.md",
    "devils_advocate": "devils_advocate_reviewer_agent.md",
    "synthesizer": "editorial_synthesizer_agent.md",
}


REVIEWER_ROLES = [
    "eic",
    "methodology",
    "domain",
    "perspective",
    "devils_advocate",
]


ROLE_LABELS = {
    "field_analyst": "领域分析 / Field Analyst",
    "eic": "主编 / Editor-in-Chief",
    "methodology": "评审人1：方法学 / Reviewer 1: Methodology",
    "domain": "评审人2：领域专家 / Reviewer 2: Domain",
    "perspective": "评审人3：跨学科视角 / Reviewer 3: Perspective",
    "devils_advocate": "反方评审 / Devil's Advocate",
    "synthesizer": "编辑综合 / Editorial Synthesizer",
}


RECOMMENDATION_ZH = {
    "Accept": "接收",
    "Minor Revision": "小修",
    "Major Revision": "大修",
    "Reject": "拒稿",
    "Panel configured": "评审团队已配置",
    "Review Complete": "评审完成",
}


SEVERITY_ZH = {
    "Critical": "致命",
    "Major": "主要",
    "Minor": "次要",
    "Observation": "观察",
}


DESCRIPTOR_ZH = {
    "Exceptional": "优秀",
    "Strong": "较强",
    "Adequate": "尚可",
    "Weak": "较弱",
    "Insufficient": "不足",
}


DIMENSION_ZH = {
    "Originality": "原创性",
    "Methodological Rigor": "方法严谨性",
    "Evidence Sufficiency": "证据充分性",
    "Argument Coherence": "论证连贯性",
    "Literature Integration": "文献整合",
    "Presentation Quality": "呈现质量",
}


@dataclass
class ManuscriptProfile:
    title: str
    abstract: str
    word_count: int
    reference_count: int
    section_names: list[str]
    language: str
    primary_discipline: str
    secondary_disciplines: list[str]
    research_paradigm: str
    methodology_type: str
    target_tier: str
    maturity: str
    signals: dict[str, Any]


@dataclass
class AgentResult:
    role: str
    label: str
    recommendation: str
    confidence: int
    summary: str
    strengths: list[str]
    weaknesses: list[dict[str, str]]
    questions: list[str]
    minor_issues: list[str]
    scores: dict[str, int]
    markdown: str


def load_agent_prompt(agent: str) -> str:
    path = AGENT_ROOT / AGENT_FILES[agent]
    return path.read_text(encoding="utf-8")


def read_text_from_upload(filename: str, payload: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return payload.decode("utf-8", errors="replace")
    if suffix == ".docx":
        return _read_docx(payload)
    if suffix == ".pdf":
        return _read_pdf(payload)
    raise ValueError("Unsupported file type. Please upload .txt, .md, .docx, or .pdf.")


def _read_docx(payload: bytes) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover - environment dependent
        raise ValueError("python-docx is required to parse .docx files.") from exc

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(payload)
        tmp_path = tmp.name
    try:
        document = docx.Document(tmp_path)
        parts: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def _read_pdf(payload: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(payload)
        tmp_path = tmp.name
    out_path = f"{tmp_path}.txt"
    try:
        subprocess.run(
            ["pdftotext", "-layout", tmp_path, out_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        return Path(out_path).read_text(encoding="utf-8", errors="replace")
    except (FileNotFoundError, subprocess.CalledProcessError) as exc:
        raise ValueError("pdftotext is required to parse .pdf files.") from exc
    finally:
        Path(tmp_path).unlink(missing_ok=True)
        Path(out_path).unlink(missing_ok=True)


def analyze_manuscript(text: str) -> ManuscriptProfile:
    cleaned = normalize_text(text)
    title = infer_title(cleaned)
    abstract = infer_abstract(cleaned)
    word_count = len(re.findall(r"\b[\w'-]+\b", cleaned))
    reference_count = count_references(cleaned)
    section_names = infer_sections(cleaned)
    language = infer_language(cleaned)
    discipline = infer_primary_discipline(cleaned)
    secondary = infer_secondary_disciplines(cleaned, discipline)
    paradigm, method_type = infer_methodology(cleaned)
    maturity = infer_maturity(cleaned, word_count, reference_count, section_names)
    tier = infer_target_tier(maturity, word_count, reference_count, section_names)
    signals = collect_signals(cleaned, section_names, reference_count)
    return ManuscriptProfile(
        title=title,
        abstract=abstract,
        word_count=word_count,
        reference_count=reference_count,
        section_names=section_names,
        language=language,
        primary_discipline=discipline,
        secondary_disciplines=secondary,
        research_paradigm=paradigm,
        methodology_type=method_type,
        target_tier=tier,
        maturity=maturity,
        signals=signals,
    )


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_title(text: str) -> str:
    for line in text.splitlines():
        candidate = line.strip(" #\t")
        if 8 <= len(candidate) <= 180 and not re.match(r"^(abstract|摘要|introduction|引言)\b", candidate, re.I):
            return candidate
    return "Untitled Manuscript"


def infer_abstract(text: str) -> str:
    match = re.search(
        r"(?:^|\n)\s*(abstract|摘要)\s*[:：]?\s*\n?(.*?)(?=\n\s*(introduction|引言|keywords|关键词|1\.|\#)|\Z)",
        text,
        re.I | re.S,
    )
    if match:
        return re.sub(r"\s+", " ", match.group(2)).strip()[:1500]
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if len(p.strip()) > 80]
    return re.sub(r"\s+", " ", paragraphs[0]).strip()[:1500] if paragraphs else ""


def count_references(text: str) -> int:
    ref_match = re.search(r"(?:^|\n)\s*(references|bibliography|参考文献)\s*\n(.*)$", text, re.I | re.S)
    ref_text = ref_match.group(2) if ref_match else text[-8000:]
    numbered = len(re.findall(r"(?:^|\n)\s*(?:\[\d+\]|\d+\.|\(\d+\))\s+", ref_text))
    author_year = len(re.findall(r"\([A-Z][A-Za-z\-]+(?: et al\.)?,?\s+(?:19|20)\d{2}\)", text))
    chinese_refs = len(re.findall(r"(?:^|\n)\s*[\u4e00-\u9fffA-Za-z].{5,80}(?:19|20)\d{2}", ref_text))
    return max(numbered, min(author_year, 80), chinese_refs if ref_match else 0)


def infer_sections(text: str) -> list[str]:
    patterns = [
        "abstract",
        "introduction",
        "literature review",
        "theoretical framework",
        "methodology",
        "methods",
        "research design",
        "results",
        "findings",
        "discussion",
        "conclusion",
        "limitations",
        "references",
        "摘要",
        "引言",
        "文献综述",
        "理论框架",
        "研究方法",
        "方法",
        "结果",
        "发现",
        "讨论",
        "结论",
        "局限",
        "参考文献",
    ]
    found: list[str] = []
    for raw in patterns:
        pat = rf"(?:^|\n)\s*(?:\d+(?:\.\d+)*[.)、]?\s*)?{re.escape(raw)}\s*(?:\n|$)"
        if re.search(pat, text, re.I):
            found.append(raw)
    return found


def infer_language(text: str) -> str:
    sample = text[:5000]
    cjk = len(re.findall(r"[\u4e00-\u9fff]", sample))
    latin = len(re.findall(r"[A-Za-z]", sample))
    return "Chinese" if cjk > latin * 0.35 else "English"


def infer_primary_discipline(text: str) -> str:
    low = text.lower()
    disciplines = {
        "education / higher education": [
            "education",
            "student",
            "teacher",
            "curriculum",
            "learning",
            "higher education",
            "university",
            "教育",
            "高校",
            "大学",
            "学生",
            "课程",
        ],
        "artificial intelligence / information science": [
            "artificial intelligence",
            "machine learning",
            "deep learning",
            "algorithm",
            "llm",
            "large language model",
            "人工智能",
            "机器学习",
            "算法",
            "大语言模型",
        ],
        "materials science": [
            "materials",
            "alloy",
            "microstructure",
            "coercivity",
            "magnet",
            "材料",
            "合金",
            "显微组织",
            "矫顽力",
            "磁",
        ],
        "medicine / health sciences": [
            "clinical",
            "patient",
            "hospital",
            "therapy",
            "diagnosis",
            "health",
            "医学",
            "临床",
            "患者",
            "诊断",
        ],
        "public policy": [
            "policy",
            "governance",
            "regulation",
            "stakeholder",
            "政策",
            "治理",
            "监管",
        ],
        "business / management": [
            "management",
            "strategy",
            "organization",
            "firm",
            "market",
            "管理",
            "战略",
            "组织",
            "市场",
        ],
        "social sciences": [
            "society",
            "social",
            "survey",
            "interview",
            "qualitative",
            "社会",
            "访谈",
            "问卷",
        ],
    }
    scores = {name: sum(low.count(term.lower()) for term in terms) for name, terms in disciplines.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "interdisciplinary / general academic"


def infer_secondary_disciplines(text: str, primary: str) -> list[str]:
    low = text.lower()
    candidates = []
    for name in [
        "education / higher education",
        "artificial intelligence / information science",
        "public policy",
        "business / management",
        "social sciences",
        "ethics / law",
        "statistics / data science",
    ]:
        if name == primary:
            continue
        key_terms = name.replace(" / ", " ").split()
        score = sum(low.count(t) for t in key_terms)
        if name == "ethics / law":
            score += sum(low.count(t) for t in ["ethics", "privacy", "fairness", "law", "伦理", "隐私", "公平", "法律"])
        if name == "statistics / data science":
            score += sum(low.count(t) for t in ["regression", "model", "dataset", "statistics", "统计", "模型", "数据"])
        if score > 0:
            candidates.append((score, name))
    candidates.sort(reverse=True)
    return [name for _, name in candidates[:3]] or ["none obvious"]


def infer_methodology(text: str) -> tuple[str, str]:
    low = text.lower()
    method_terms = {
        ("Quantitative Research", "Survey / Questionnaire"): ["survey", "questionnaire", "likert", "问卷", "量表"],
        ("Quantitative Research", "Statistical Modeling / Machine Learning"): [
            "regression",
            "anova",
            "statistical",
            "machine learning",
            "classification",
            "回归",
            "统计",
            "机器学习",
        ],
        ("Qualitative Research", "Interview / Case Study"): ["interview", "case study", "thematic analysis", "访谈", "案例", "主题分析"],
        ("Mixed Methods", "Mixed Methods Design"): ["mixed methods", "qualitative and quantitative", "混合方法"],
        ("Literature Review / Meta-analysis", "Systematic Review / Scoping Review"): [
            "systematic review",
            "scoping review",
            "meta-analysis",
            "prisma",
            "文献综述",
            "元分析",
            "系统综述",
        ],
        ("Theoretical/Conceptual Analysis", "Conceptual / Theoretical Analysis"): [
            "conceptual",
            "theoretical",
            "framework",
            "概念",
            "理论",
        ],
        ("Experimental Research", "Experimental / Quasi-experimental"): [
            "experiment",
            "randomized",
            "control group",
            "pretest",
            "posttest",
            "实验",
            "对照组",
        ],
    }
    scored = []
    for pair, terms in method_terms.items():
        scored.append((sum(low.count(term) for term in terms), pair))
    scored.sort(reverse=True, key=lambda item: item[0])
    if scored and scored[0][0] > 0:
        return scored[0][1]
    return "Unclear / Not explicitly stated", "Unclear"


def infer_maturity(text: str, word_count: int, reference_count: int, sections: list[str]) -> str:
    has_core = sum(bool(any(token in sec for sec in sections)) for token in ["introduction", "method", "results", "discussion", "conclusion"])
    has_cn_core = sum(bool(any(token in sec for sec in sections)) for token in ["引言", "方法", "结果", "讨论", "结论"])
    core_count = max(has_core, has_cn_core)
    if word_count >= 4500 and reference_count >= 20 and core_count >= 4:
        return "Pre-submission"
    if word_count >= 2500 and reference_count >= 8 and core_count >= 3:
        return "Revised draft"
    return "First draft"


def infer_target_tier(maturity: str, word_count: int, reference_count: int, sections: list[str]) -> str:
    if maturity == "Pre-submission" and reference_count >= 35 and word_count >= 6500:
        return "Q1/Q2 candidate, depending on novelty and evidence strength"
    if maturity in {"Pre-submission", "Revised draft"}:
        return "Q2/Q3 candidate"
    return "Q3/Q4 developmental target"


def collect_signals(text: str, sections: list[str], reference_count: int) -> dict[str, Any]:
    low = text.lower()
    return {
        "has_methods": any(sec in sections for sec in ["methodology", "methods", "research design", "研究方法", "方法"]),
        "has_results": any(sec in sections for sec in ["results", "findings", "结果", "发现"]),
        "has_limitations": any(sec in sections for sec in ["limitations", "局限"]) or "limitation" in low or "局限" in text,
        "has_references": reference_count > 0,
        "has_data_availability": "data availability" in low or "数据可得" in text,
        "has_ethics": any(term in low for term in ["ethics", "irb", "informed consent"]) or any(term in text for term in ["伦理", "知情同意"]),
        "has_effect_size": any(term in low for term in ["effect size", "cohen", "eta squared", "confidence interval", "95% ci"]),
        "causal_language": any(term in low for term in ["cause", "causal", "impact", "effect"]) or any(term in text for term in ["因果", "影响", "效应"]),
        "overclaim_terms": any(term in low for term in ["prove", "definitively", "always", "all cases"]) or any(term in text for term in ["证明", "必然", "所有", "完全"]),
    }


def build_field_report(profile: ManuscriptProfile) -> str:
    cards = reviewer_configuration_cards(profile)
    sections = ", ".join(profile.section_names) if profile.section_names else "未清晰标出 / not clearly marked"
    secondaries = ", ".join(profile.secondary_disciplines)
    return f"""# 领域分析报告 / Field Analysis Report

## 稿件基本信息 / Paper Basic Information
- **题目 / Title**: {profile.title}
- **摘要长度 / Abstract length**: {len(profile.abstract.split()) if profile.abstract else 0} words
- **全文长度 / Full text length**: approximately {profile.word_count} words
- **参考文献数量 / Number of references**: approximately {profile.reference_count}
- **检测到的章节 / Detected sections**: {sections}

## 领域分析 / Field Analysis

| 维度 / Dimension | 分析结果 / Analysis Result |
|-----------|----------------|
| 主要学科 / Primary Discipline | {profile.primary_discipline} |
| 相关学科 / Secondary Disciplines | {secondaries} |
| 研究范式 / Research Paradigm | {profile.research_paradigm} |
| 方法类型 / Methodology Type | {profile.methodology_type} |
| 目标期刊层级 / Target Journal Tier | {profile.target_tier} |
| 稿件成熟度 / Paper Maturity | {profile.maturity} |

## 评审团队配置 / Reviewer Configuration Cards

{cards}

## 评审策略建议 / Review Strategy Recommendations
- 五个 Phase 1 评审角色应保持独立：主编关注期刊匹配与贡献，R1 关注方法，R2 关注文献与领域准确性，R3 关注跨学科和实践盲点，反方评审负责压力测试核心论证。
- 本轮应作为投稿前诊断使用。系统只生成评审意见，不会改写或编辑原稿。
"""


def reviewer_configuration_cards(profile: ManuscriptProfile) -> str:
    discipline = profile.primary_discipline
    method = profile.methodology_type
    secondary = profile.secondary_disciplines[0] if profile.secondary_disciplines else "an adjacent field"
    return f"""### 评审配置卡 #1 / Reviewer Configuration Card #1

**角色 / Role**: 主编 / EIC
**身份描述 / Identity Description**: 国际 {discipline} 期刊的资深编辑，重点判断期刊匹配度、原创性、学术贡献和读者价值。
**评审重点 / Review Focus**:
  1. 稿件是否对 {discipline} 提供清晰贡献。
  2. 题目、摘要、引言和结论是否一致。
  3. 当前成熟度是否支撑目标层级：{profile.target_tier}。
**特别关注 / Will particularly care about**: 研究野心、证据强度和贡献表述之间是否匹配。
**可能盲点 / Possible blind spots**: 不会深入检查具体方法技术细节。

### 评审配置卡 #2 / Reviewer Configuration Card #2

**角色 / Role**: 评审人1：方法学 / Peer Reviewer 1 — Methodology
**身份描述 / Identity Description**: {method} 专家，评估研究设计有效性、分析选择、报告透明度和可复现性。
**评审重点 / Review Focus**:
  1. 方法是否能够回答研究问题。
  2. 抽样、数据收集和分析是否交代充分。
  3. 结论是否严格限定在证据范围内。
**特别关注 / Will particularly care about**: 设计细节缺失、因果推断越界和可复现性不足。
**可能盲点 / Possible blind spots**: 可能低估更广泛的理论贡献。

### 评审配置卡 #3 / Reviewer Configuration Card #3

**角色 / Role**: 评审人2：领域专家 / Peer Reviewer 2 — Domain
**身份描述 / Identity Description**: {discipline} 资深学者，关注文献覆盖、概念精确性和领域贡献。
**评审重点 / Review Focus**:
  1. 是否覆盖关键文献和核心争论。
  2. 理论框架是否准确，且真正用于分析。
  3. 所声称的贡献是否确有新意。
**特别关注 / Will particularly care about**: 基础文献或近年重要研究缺失、概念混用。
**可能盲点 / Possible blind spots**: 可能对跨学科主张较为保守。

### 评审配置卡 #4 / Reviewer Configuration Card #4

**角色 / Role**: 评审人3：跨学科/实践视角 / Peer Reviewer 3 — Cross-disciplinary/Practical
**身份描述 / Identity Description**: 来自 {secondary} 的评审人，从外部视角审视前提、可行性、伦理和更广泛影响。
**评审重点 / Review Focus**:
  1. 是否处理利益相关者视角和实施约束。
  2. 关键假设能否跨情境成立。
  3. 是否可引入相邻领域概念来增强论文。
**特别关注 / Will particularly care about**: 实践价值和被遗漏的视角。
**可能盲点 / Possible blind spots**: 可能不完全遵循主学科的发表惯例。

### 评审配置卡 #5 / Reviewer Configuration Card #5

**角色 / Role**: 反方评审 / Devil's Advocate
**身份描述 / Identity Description**: 对抗式逻辑评审，压力测试中心论点、证据链和最强反驳。
**评审重点 / Review Focus**:
  1. 找出对论文最有力的反驳。
  2. 识别逻辑缺口、过度概括和缺乏支撑的核心主张。
  3. 标出会阻止接收的关键问题。
**特别关注 / Will particularly care about**: 主要结论是否真正从证据中推出。
**可能盲点 / Possible blind spots**: 有意偏向挑战和压力测试，而不是平衡评价。"""


def run_review(text: str, options: dict[str, Any] | None = None, progress: Any = None) -> dict[str, Any]:
    options = options or {}
    profile = analyze_manuscript(text)
    provider = options.get("provider") or os.environ.get("AI_REVIEWER_PROVIDER", "heuristic")
    mode = options.get("mode", "full")
    model = options.get("model") or os.environ.get("OPENAI_MODEL", "")
    base_url = options.get("base_url") or os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1")
    api_key = options.get("api_key") or os.environ.get("OPENAI_API_KEY", "")
    disable_proxy = bool(options.get("disable_proxy", False))
    verify_ssl = bool(options.get("verify_ssl", True))
    ca_bundle = options.get("ca_bundle") or os.environ.get("SSL_CERT_FILE", "")

    if provider == "openai":
        if not api_key:
            raise ValueError("后端没有配置 API Key。请在 app/config.json 中配置 api_key，或设置环境变量 OPENAI_API_KEY。")
        return run_llm_review(
            text,
            profile,
            mode=mode,
            model=model,
            base_url=base_url,
            api_key=api_key,
            disable_proxy=disable_proxy,
            verify_ssl=verify_ssl,
            ca_bundle=ca_bundle,
            progress=progress,
        )
    return run_heuristic_review(text, profile, mode=mode, progress=progress)


def emit_progress(progress: Any, agent_id: str, status: str, **extra: Any) -> None:
    if progress:
        progress(agent_id, status, **extra)


def run_heuristic_review(text: str, profile: ManuscriptProfile, mode: str = "full", progress: Any = None) -> dict[str, Any]:
    emit_progress(progress, "field_analyst", "running", message="正在分析领域和配置评审团队")
    field_report = build_field_report(profile)
    emit_progress(progress, "field_analyst", "complete", recommendation=rec_label("Panel configured"), confidence=4)
    reviewer_results = []
    for role in REVIEWER_ROLES:
        emit_progress(progress, role, "running", message=f"{ROLE_LABELS[role]} 正在生成评审意见")
        result = heuristic_agent_review(role, profile)
        reviewer_results.append(result)
        emit_progress(progress, role, "complete", recommendation=rec_label(result.recommendation), confidence=result.confidence)
    emit_progress(progress, "synthesizer", "running", message="正在综合多位评审意见")
    synthesis = synthesize_reviews(profile, reviewer_results)
    emit_progress(progress, "synthesizer", "complete", recommendation=rec_label(synthesis["decision"]), confidence=synthesis["confidence"])
    return {
        "provider": "heuristic",
        "mode": mode,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "profile": profile_to_dict(profile),
        "agents": [
            {
                "id": "field_analyst",
                "label": ROLE_LABELS["field_analyst"],
                "status": "complete",
                "recommendation": rec_label("Panel configured"),
                "confidence": 4,
                "markdown": field_report,
            },
            *[
                {
                    "id": result.role,
                    "label": result.label,
                    "status": "complete",
                    "recommendation": rec_label(result.recommendation),
                    "confidence": result.confidence,
                    "markdown": result.markdown,
                    "scores": result.scores,
                }
                for result in reviewer_results
            ],
            {
                "id": "synthesizer",
                "label": ROLE_LABELS["synthesizer"],
                "status": "complete",
                "recommendation": rec_label(synthesis["decision"]),
                "confidence": synthesis["confidence"],
                "markdown": synthesis["markdown"],
            },
        ],
        "summary": synthesis,
    }


def heuristic_agent_review(role: str, profile: ManuscriptProfile) -> AgentResult:
    if role == "eic":
        return eic_review(profile)
    if role == "methodology":
        return methodology_review(profile)
    if role == "domain":
        return domain_review(profile)
    if role == "perspective":
        return perspective_review(profile)
    if role == "devils_advocate":
        return devils_advocate_review(profile)
    raise ValueError(f"Unknown role: {role}")


def base_scores(profile: ManuscriptProfile) -> dict[str, int]:
    score = {
        "Originality": 68,
        "Methodological Rigor": 68,
        "Evidence Sufficiency": 68,
        "Argument Coherence": 70,
        "Literature Integration": 68,
        "Presentation Quality": 70,
    }
    if profile.maturity == "Pre-submission":
        for key in score:
            score[key] += 8
    elif profile.maturity == "First draft":
        for key in score:
            score[key] -= 12
    if profile.reference_count < 10:
        score["Literature Integration"] -= 15
        score["Evidence Sufficiency"] -= 5
    elif profile.reference_count > 30:
        score["Literature Integration"] += 7
    if not profile.signals["has_methods"]:
        score["Methodological Rigor"] -= 18
        score["Evidence Sufficiency"] -= 8
    if not profile.signals["has_results"]:
        score["Evidence Sufficiency"] -= 12
    if not profile.signals["has_limitations"]:
        score["Argument Coherence"] -= 6
    if profile.signals["causal_language"] and profile.methodology_type not in {"Experimental / Quasi-experimental", "Statistical Modeling / Machine Learning"}:
        score["Argument Coherence"] -= 8
        score["Methodological Rigor"] -= 6
    if profile.signals["overclaim_terms"]:
        score["Argument Coherence"] -= 8
    return {key: max(30, min(92, value)) for key, value in score.items()}


def recommendation_from_scores(scores: dict[str, int], critical: bool = False) -> str:
    if critical:
        return "Reject"
    mean = statistics.mean(scores.values())
    floor = min(scores.values())
    if floor < 45 or mean < 55:
        return "Reject"
    if floor < 60 or mean < 68:
        return "Major Revision"
    if floor < 75 or mean < 80:
        return "Minor Revision"
    return "Accept"


def descriptor(score: int) -> str:
    if score >= 85:
        return "Exceptional"
    if score >= 75:
        return "Strong"
    if score >= 65:
        return "Adequate"
    if score >= 50:
        return "Weak"
    return "Insufficient"


def rec_label(value: str) -> str:
    zh = RECOMMENDATION_ZH.get(value)
    return f"{zh} / {value}" if zh else value


def severity_label(value: str) -> str:
    zh = SEVERITY_ZH.get(value)
    return f"{zh} / {value}" if zh else value


def descriptor_label(value: str) -> str:
    zh = DESCRIPTOR_ZH.get(value)
    return f"{zh} / {value}" if zh else value


def dimension_label(value: str) -> str:
    zh = DIMENSION_ZH.get(value)
    return f"{zh} / {value}" if zh else value


def confidence_for(profile: ManuscriptProfile, role: str) -> int:
    if role == "domain" and profile.primary_discipline == "interdisciplinary / general academic":
        return 3
    if role == "methodology" and profile.methodology_type == "Unclear":
        return 3
    if profile.word_count < 1200:
        return 3
    return 4


def eic_review(profile: ManuscriptProfile) -> AgentResult:
    scores = base_scores(profile)
    scores["Originality"] += 2 if profile.secondary_disciplines and profile.secondary_disciplines[0] != "none obvious" else -4
    scores["Presentation Quality"] += 3 if len(profile.section_names) >= 6 else -8
    scores = clamp_scores(scores)
    rec = recommendation_from_scores(scores)
    strengths = [
        f"稿件能够归入 {profile.primary_discipline}，这为投稿定位和编辑判断提供了基本框架。",
        f"系统检测到的研究范式为 {profile.research_paradigm}，方法类型为 {profile.methodology_type}，说明稿件已有可供评审的基本研究形态。",
        "稿件主题具备进行投稿前发展性评审的基础，后续重点应放在贡献定位和证据链打磨上。",
    ]
    weaknesses = []
    if profile.maturity == "First draft":
        weaknesses.append(issue("稿件成熟度偏低", "从结构、参考文献数量或证据呈现来看，目前更接近初稿，而不是可直接投稿的期刊论文。", "先稳定完整的 IMRaD 或本领域适用结构，再进入投稿语言和格式层面的打磨。", "Major"))
    if len(profile.section_names) < 5:
        weaknesses.append(issue("章节结构不够清晰", "若干期刊论文常见章节没有被清楚标出，读者难以快速判断论证路径。", "增加明确的小节标题，并让引言、方法、结果、讨论、结论之间形成可追踪的逻辑链。", "Major"))
    if profile.reference_count < 15:
        weaknesses.append(issue("贡献主张支撑不足", "参考文献基础偏薄，尚不足以支撑面向期刊读者的学术贡献声明。", "补充本领域奠基性研究、近三年关键研究和最接近的竞争性文献，并进行综合比较。", "Major"))
    while len(weaknesses) < 3:
        weaknesses.append(issue("期刊匹配需要更精准", "目前尚不能清楚看出目标读者、目标期刊层级和论文贡献之间的对应关系。", "明确目标期刊类型，并据此调整研究问题范围、贡献表述和讨论重点。", "Minor"))
    questions = [
        "这篇稿件准备投向哪一类期刊或哪一个期刊层级？",
        "如果只能保留一个核心贡献，作者希望读者记住的贡献是什么？它与相邻研究相比具体新在哪里？",
    ]
    minor = ["投稿前检查题目、摘要和结论是否使用同一套贡献表述。", "统一章节标题层级，并对齐目标期刊格式。"]
    return make_result("eic", profile, rec, confidence_for(profile, "eic"), strengths, weaknesses, questions, minor, scores)


def methodology_review(profile: ManuscriptProfile) -> AgentResult:
    scores = base_scores(profile)
    if not profile.signals["has_methods"]:
        scores["Methodological Rigor"] -= 12
    if profile.signals["causal_language"] and "Experimental" not in profile.methodology_type:
        scores["Argument Coherence"] -= 7
    if profile.signals["has_data_availability"]:
        scores["Methodological Rigor"] += 4
    if profile.signals["has_effect_size"]:
        scores["Evidence Sufficiency"] += 4
    scores = clamp_scores(scores)
    rec = recommendation_from_scores(scores)
    strengths = [
        f"稿件已有足够文本线索，可初步判断方法类型为 {profile.methodology_type}。",
        "研究设计、证据和结论之间的对应关系已经可以被评审定位，说明后续修改有明确抓手。",
        "只要补足程序、样本、分析假设和有效性检查，稿件的方法透明度可以明显提升。",
    ]
    weaknesses = []
    if not profile.signals["has_methods"]:
        weaknesses.append(issue("方法部分不清晰", "稿件中没有检测到明确的方法或研究设计章节。", "增加独立方法部分，交代研究设计、样本/语料、数据收集、分析流程和有效性检查。", "Critical"))
    if not profile.signals["has_results"]:
        weaknesses.append(issue("结果证据难以定位", "没有检测到独立的结果或发现章节，评审人难以判断结论来自哪些证据。", "将结果与讨论分开，并把每一项结果对应到具体研究问题。", "Major"))
    if profile.signals["causal_language"] and "Experimental" not in profile.methodology_type:
        weaknesses.append(issue("因果表述可能超过研究设计承载能力", "稿件使用了 impact/effect 等影响或效应语言，但没有明确实验或准实验设计信号。", "改用相关性或关联性表述；如果必须保留因果主张，需要补充识别策略、稳健性检查和混杂因素讨论。", "Major"))
    if not profile.signals["has_data_availability"]:
        weaknesses.append(issue("可复现材料不可见", "没有检测到数据可得性、代码可得性或复现说明。", "增加 data/code availability statement；若无法公开，应说明限制原因和可替代验证方式。", "Minor"))
    while len(weaknesses) < 3:
        weaknesses.append(issue("分析假设报告不足", "当前文本没有充分显示统计或编码分析的前提假设检查。", "按方法类型补充假设检验、缺失数据处理、敏感性分析或编码一致性说明。", "Minor"))
    questions = [
        "样本、语料或案例是如何选择的？纳入和排除标准是什么？",
        "哪些有效性检查、稳健性检查或编码一致性证据支持主要发现？",
        "数据、代码、访谈提纲或编码方案是否可以公开？如果不能，原因是什么？",
    ]
    minor = ["为量表、数据集、访谈提纲或实验协议补充页码/行号定位。", "说明是否保留了不显著结果、负面结果或与预期相反的发现。"]
    return make_result("methodology", profile, rec, confidence_for(profile, "methodology"), strengths, weaknesses[:5], questions, minor, scores)


def domain_review(profile: ManuscriptProfile) -> AgentResult:
    scores = base_scores(profile)
    if profile.reference_count < 15:
        scores["Literature Integration"] -= 14
    if profile.reference_count > 30:
        scores["Literature Integration"] += 6
    if profile.primary_discipline == "interdisciplinary / general academic":
        scores["Originality"] -= 6
    scores = clamp_scores(scores)
    rec = recommendation_from_scores(scores)
    strengths = [
        f"选题可以放在 {profile.primary_discipline} 中讨论，这使论文有机会形成面向特定领域的贡献主张。",
        "稿件已经出现若干领域线索，便于后续补充和整合相邻文献。",
        "如果进一步明确核心概念和理论边界，论文贡献会更容易被领域读者识别。",
    ]
    weaknesses = []
    if profile.reference_count < 15:
        weaknesses.append(issue("文献基础偏薄", f"系统仅检测到约 {profile.reference_count} 条参考文献。", "补充奠基性研究、近年高影响力研究，以及与本文贡献最接近的直接竞争文献。", "Major"))
    weaknesses.append(issue("研究空白需要更清楚地区分", "目前的稿件画像还不能证明该研究空白已经与相邻研究清楚区分。", "明确已有研究已经解释了什么、仍未解决什么，以及本文如何改变现有知识状态。", "Major" if scores["Originality"] < 65 else "Minor"))
    weaknesses.append(issue("核心概念定义需要审计", "核心术语可能在全文中没有保持一致定义。", "增加简短概念表或定义段落，说明关键构念、操作化方式和使用边界。", "Minor"))
    questions = [
        f"在 {profile.primary_discipline} 中，与本文最接近的三篇研究是什么？本文与它们的具体差异在哪里？",
        "近三年是否有会改变本文框架、研究空白或结论解释的重要研究？",
    ]
    minor = ["重要理论尽量引用原始出处，而不是只引用综述或二手转述。", "引入核心概念时避免依赖二手引用。"]
    return make_result("domain", profile, rec, confidence_for(profile, "domain"), strengths, weaknesses, questions, minor, scores)


def perspective_review(profile: ManuscriptProfile) -> AgentResult:
    scores = base_scores(profile)
    if not profile.signals["has_limitations"]:
        scores["Argument Coherence"] -= 7
    if not profile.signals["has_ethics"] and any(term in profile.primary_discipline for term in ["education", "medicine", "social", "artificial intelligence"]):
        scores["Evidence Sufficiency"] -= 4
    scores = clamp_scores(scores)
    rec = recommendation_from_scores(scores)
    secondary = ", ".join(profile.secondary_disciplines)
    strengths = [
        f"稿件与 {secondary} 存在可发展的跨学科连接。",
        "如果作者说明哪些主体能够基于研究发现采取行动，论文可以形成更有价值的 broader implications。",
        "当前稿件适合补充利益相关者、实施条件和潜在风险分析。",
    ]
    weaknesses = []
    if not profile.signals["has_limitations"]:
        weaknesses.append(issue("局限性与可迁移性讨论不足", "没有检测到清晰的局限性章节。", "补充研究范围、情境边界、数据限制和结论可迁移性的讨论。", "Major"))
    if not profile.signals["has_ethics"] and any(term in profile.primary_discipline for term in ["education", "medicine", "social", "artificial intelligence"]):
        weaknesses.append(issue("伦理与利益相关者影响需要补充", "该主题可能影响真实人群或组织，但目前没有看到伦理、隐私、公平或利益相关者保护讨论。", "根据主题补充隐私、公平、知情同意、权力不对称或受影响群体分析。", "Major"))
    weaknesses.append(issue("实践落地路径不明确", "研究发现的实践后果还没有被操作化。", "说明谁应在什么约束下采取什么行动，以及可能产生哪些风险或副作用。", "Minor"))
    questions = [
        "论文建议最直接影响哪一类利益相关者？他们的视角是否已经在稿件中呈现？",
        "如果要把结论迁移到研究场景之外，需要满足哪些情境条件？",
    ]
    minor = ["增加一小段关于 unintended consequences 的讨论。", "区分学术启示和实践建议，不要把两者混写。"]
    return make_result("perspective", profile, rec, confidence_for(profile, "perspective"), strengths, weaknesses, questions, minor, scores)


def devils_advocate_review(profile: ManuscriptProfile) -> AgentResult:
    scores = base_scores(profile)
    critical = False
    issues = []
    if not profile.signals["has_methods"] and not profile.signals["has_results"]:
        critical = True
        issues.append(issue("证据链不可评审", "没有检测到清晰的方法部分，也没有检测到清晰的结果部分。", "在提出期刊级贡献主张前，先完整呈现研究设计、证据来源和结果链条。", "Critical"))
    if profile.signals["causal_language"] and "Experimental" not in profile.methodology_type:
        issues.append(issue("结论可能不能从研究设计中推出", "稿件使用 impact/effect 等因果或效应语言，但没有明确因果设计信号。", "重写为更谨慎的关联性表述，或补充因果识别与混杂控制。", "Major"))
    if profile.reference_count < 8:
        issues.append(issue("反向文献可能缺失", "参考文献数量过少，无法证明作者已经处理了相反发现或边界条件。", "补充具有代表性的反向证据、失败案例和边界条件研究。", "Major"))
    if profile.signals["overclaim_terms"]:
        issues.append(issue("存在过度概括风险", "稿件中检测到绝对化语言。", "将普遍性断言改为受证据、样本和情境约束的有限主张。", "Major"))
    while len(issues) < 3:
        issues.append(issue("最强替代解释没有被正面处理", "稿件没有清楚呈现能够挑战作者解释的 rival explanation。", "增加替代解释段落，并说明为什么本文解释比替代解释更有说服力。", "Major"))

    rec = "Reject" if critical else ("Major Revision" if any(i["severity"] == "Major" for i in issues) else "Minor Revision")
    strengths = [
        "稿件已经有可以被压力测试的中心议题。",
        "若能更清楚限定范围并加强证据纪律，若干潜在漏洞是可以修复的。",
        "更强的局限性和反驳处理部分会显著提升真实评审人的信任度。",
    ]
    questions = [
        "对于主要发现，最有力的替代解释是什么？为什么它不如作者解释有说服力？",
        "如果移除最弱的一项证据，论文中哪一个核心主张会首先失败？",
    ]
    minor = ["除非证据确实具有普遍性，否则避免使用绝对化措辞。", "超出直接数据范围的判断应使用更谨慎的情态表达。"]
    result = make_result("devils_advocate", profile, rec, 4 if profile.word_count > 1200 else 3, strengths, issues, questions, minor, scores)
    result.markdown = devil_markdown(profile, result, issues)
    return result


def issue(title: str, problem: str, suggestion: str, severity: str) -> dict[str, str]:
    return {"title": title, "problem": problem, "suggestion": suggestion, "severity": severity}


def clamp_scores(scores: dict[str, int]) -> dict[str, int]:
    return {key: max(30, min(94, int(value))) for key, value in scores.items()}


def make_result(
    role: str,
    profile: ManuscriptProfile,
    recommendation: str,
    confidence: int,
    strengths: list[str],
    weaknesses: list[dict[str, str]],
    questions: list[str],
    minor_issues: list[str],
    scores: dict[str, int],
) -> AgentResult:
    result = AgentResult(
        role=role,
        label=ROLE_LABELS[role],
        recommendation=recommendation,
        confidence=confidence,
        summary=summary_for(role, profile, recommendation),
        strengths=strengths,
        weaknesses=weaknesses,
        questions=questions,
        minor_issues=minor_issues,
        scores=scores,
        markdown="",
    )
    result.markdown = peer_markdown(profile, result)
    return result


def summary_for(role: str, profile: ManuscriptProfile, recommendation: str) -> str:
    focus = {
        "eic": "期刊匹配、学术贡献和稿件成熟度",
        "methodology": "研究设计、证据质量和可复现性",
        "domain": "文献覆盖、概念精确性和领域贡献",
        "perspective": "跨学科价值、利益相关者影响和实践可迁移性",
        "devils_advocate": "逻辑漏洞、反驳压力和过度主张风险",
    }[role]
    return (
        f"从{focus}角度看，本稿当前更适合被判断为 **{rec_label(recommendation)}**。"
        f"系统将其定位在 {profile.primary_discipline}，成熟度为 {profile.maturity}，"
        f"全文约 {profile.word_count} words，检测到约 {profile.reference_count} 条参考文献。"
        "主要修改方向是进一步明确论文贡献、证据链条和结论边界。"
    )


def peer_markdown(profile: ManuscriptProfile, result: AgentResult) -> str:
    strengths = "\n".join(f"### S{i + 1}: {item.split('.')[0]}\n{item}" for i, item in enumerate(result.strengths))
    weaknesses = "\n\n".join(
        f"### W{i + 1}: {item['title']}\n"
        f"**问题 / Problem**: {item['problem']}\n"
        f"**重要性 / Why it matters**: 这个问题会降低评审人对论文贡献、证据支撑或投稿成熟度的信心。\n"
        f"**修改建议 / Suggestion**: {item['suggestion']}\n"
        f"**严重程度 / Severity**: {severity_label(item['severity'])}"
        for i, item in enumerate(result.weaknesses)
    )
    questions = "\n".join(f"{i + 1}. {q}" for i, q in enumerate(result.questions))
    minor = "\n".join(f"- {m}" for m in result.minor_issues)
    score_rows = "\n".join(
        f"| {dimension_label(name)} | {score} | {descriptor_label(descriptor(score))} | 本地 MVP 规则估计 / Rule-based MVP estimate |" for name, score in result.scores.items()
    )
    return f"""# 同行评审报告 / Peer Review Report

## 稿件信息 / Manuscript Information
- **题目 / Title**: {profile.title}
- **评审日期 / Review Date**: {datetime.now().date().isoformat()}
- **评审轮次 / Review Round**: 投稿前诊断 / Pre-submission diagnostic

## 评审人信息 / Reviewer Information

### 评审角色 / Reviewer Role
{result.label}

### 评审身份 / Reviewer Identity
{reviewer_identity(result.role, profile)}

### 评审重点 / Review Focus
{reviewer_focus(result.role)}

## 总体评价 / Overall Assessment

### 建议 / Recommendation
**{rec_label(result.recommendation)}**

### 置信度 / Confidence Score
**{result.confidence}/5**

置信度说明：这是该评审角色对“自己这份判断可靠程度”的自评，主要受稿件信息完整度、角色专业匹配度和证据可见性影响。它不是论文被接收的概率，也不是模型准确率。

### 摘要评价 / Summary Assessment
{result.summary}

## 优点 / Strengths

{strengths}

## 主要问题 / Weaknesses

{weaknesses}

## 分章节意见 / Detailed Comments

### 题目与摘要 / Title & Abstract
- 检查题目、摘要和结论是否使用一致的贡献表述。
- 摘要应交代研究问题、方法、核心发现和贡献，同时避免过度声称。

### 引言 / Introduction
- 引言需要更清楚呈现研究空白，并说明本文与最接近研究的区别。

### 文献综述 / 理论框架 / Literature Review / Theoretical Framework
- 文献综述应综合争论和研究脉络，而不是罗列引用。
- 理论概念应先定义，再进入分析使用。

### 方法 / 研究设计 / Methodology / Research Design
- 方法部分应交代设计、抽样、数据收集、分析方法、有效性检查和可复现性限制。

### 结果 / 发现 / Results / Findings
- 结果应与解释分开，并逐项对应研究问题。

### 讨论 / Discussion
- 讨论部分应说明理论/实践启示、局限性和替代解释。

### 结论 / Conclusion
- 结论应限定在证据能够支持的范围内，避免普遍化或绝对化表述。

## 给作者的问题 / Questions for Authors

{questions}

## 次要问题 / Minor Issues

{minor}

## 维度评分 / Dimension Scores

| 维度 / Dimension | 分数 / Score (0-100) | 描述 / Descriptor | 说明 / Notes |
|-----------|--------------|------------|-------|
{score_rows}
"""


def devil_markdown(profile: ManuscriptProfile, result: AgentResult, issues: list[dict[str, str]]) -> str:
    grouped = {"Critical": [], "Major": [], "Minor": [], "Observation": []}
    for item in issues:
        grouped.setdefault(item["severity"], []).append(item)

    def table(items: list[dict[str, str]]) -> str:
        if not items:
            return "| # | 维度 / Dimension | 问题描述 / Issue Description | 位置 / Location |\n|---|-----------|-------------------|----------|\n| - | - | 本地 MVP 规则未检测到该级别问题 | - |"
        rows = ["| # | 维度 / Dimension | 问题描述 / Issue Description | 位置 / Location |", "|---|-----------|-------------------|----------|"]
        for i, item in enumerate(items, 1):
            rows.append(f"| {i} | {item['title']} | {item['problem']} 修改建议：{item['suggestion']} | 稿件层面信号 |")
        return "\n".join(rows)

    return f"""# 反方评审 / Devil's Advocate Review

## 稿件信息 / Manuscript Information
- **题目 / Title**: {profile.title}
- **评审日期 / Review Date**: {datetime.now().date().isoformat()}

## 最强反驳 / Strongest Counter-Argument
最强的怀疑性解读是：稿件当前的研究野心可能超过了可见证据链所能支撑的范围。严格评审人可能会认为，论文选题有价值，但当前版本尚未通过足够透明的方法说明、文献定位和结论边界，排除更简单的替代解释。如果论文核心贡献依赖 impact、effect、novelty 或 practical transferability，作者需要证明这些主张确实来自研究设计和证据，而不只是来自作者偏好的解释路径。

## 问题列表 / Issue List

### 致命 / CRITICAL
{table(grouped.get("Critical", []))}

### 主要 / MAJOR
{table(grouped.get("Major", []))}

### 次要 / MINOR
{table(grouped.get("Minor", []))}

## 被忽略的替代解释或路径 / Ignored Alternative Explanations/Paths
- 稿件应明确指出最能挑战主要发现或核心论证的 rival explanation。
- 如果是概念论文，应使用反例测试框架边界。
- 如果是实证论文，应区分“证据支持的发现”和“合理但仍需限定的解释”。

## 缺失的利益相关者视角 / Missing Stakeholder Perspectives
- 补充最受论文主张或建议影响的人群视角。
- 说明谁受益、谁承担成本，以及谁可能有合理反对意见。

## 压力测试结论 / Overall Stress-Test Recommendation
**{rec_label(result.recommendation)}**。这是反方压力测试结论，不等同于平衡的最终发表裁决。
"""


def reviewer_identity(role: str, profile: ManuscriptProfile) -> str:
    identities = {
        "eic": f"国际 {profile.primary_discipline} 期刊资深编辑 / Senior editor of an international {profile.primary_discipline} journal.",
        "methodology": f"{profile.methodology_type} 研究设计专家 / Research design specialist in {profile.methodology_type}.",
        "domain": f"{profile.primary_discipline} 资深学者，关注领域定位和文献准确性 / Senior domain scholar.",
        "perspective": f"跨学科评审人，带入 {', '.join(profile.secondary_disciplines)} 视角 / Cross-disciplinary reviewer.",
        "devils_advocate": "对抗式逻辑评审，关注最强反驳 / Adversarial logic reviewer.",
    }
    return identities[role]


def reviewer_focus(role: str) -> str:
    return {
        "eic": "期刊匹配、原创性、重要性和全文一致性 / Journal fit, originality, significance, and coherence.",
        "methodology": "研究设计、抽样、分析、有效性和可复现性 / Research design, sampling, analysis, validity, and reproducibility.",
        "domain": "文献覆盖、理论框架、术语和领域贡献 / Literature, theory, terminology, and contribution.",
        "perspective": "前提假设、实践可行性、利益相关者影响和跨学科迁移 / Assumptions, feasibility, stakeholders, and transfer.",
        "devils_advocate": "核心论证脆弱点、逻辑缺口、反证和过度概括 / Argument vulnerability and counter-evidence.",
    }[role]


def synthesize_reviews(profile: ManuscriptProfile, reviews: list[AgentResult]) -> dict[str, Any]:
    rec_order = {"Accept": 0, "Minor Revision": 1, "Major Revision": 2, "Reject": 3}
    has_critical_da = any(
        review.role == "devils_advocate" and any(w["severity"] == "Critical" for w in review.weaknesses)
        for review in reviews
    )
    max_rec = max(reviews, key=lambda r: rec_order.get(r.recommendation, 2)).recommendation
    decision = "Major Revision" if max_rec == "Reject" and not has_critical_da else max_rec
    if has_critical_da:
        decision = "Reject"
    confidence = round(statistics.mean([r.confidence for r in reviews]))
    priority_1 = []
    priority_2 = []
    priority_3 = []
    for review in reviews:
        for weakness in review.weaknesses:
            line = f"**{weakness['title']}** ({review.label}): {weakness['suggestion']}"
            if weakness["severity"] == "Critical":
                priority_1.append(line)
            elif weakness["severity"] == "Major":
                priority_1.append(line)
            else:
                priority_2.append(line)
        for minor in review.minor_issues:
            priority_3.append(f"{review.label}: {minor}")

    inventory_rows = "\n".join(
        f"| {r.label} | {rec_label(r.recommendation)} | {r.confidence}/5 | {', '.join(w['title'] for w in r.weaknesses[:2])} |"
        for r in reviews
    )
    p1 = "\n".join(f"- {item}" for item in priority_1[:10]) or "- 未识别到必须修改项。"
    p2 = "\n".join(f"- {item}" for item in priority_2[:10]) or "- 未识别到建议修改项。"
    p3 = "\n".join(f"- {item}" for item in priority_3[:10]) or "- 未识别到次要修改项。"
    markdown = f"""# 编辑决策包 / Editorial Decision Package

## 第一部分：编辑决策信 / Part 1: Editorial Decision Letter

作者您好，

感谢您提交题为 "{profile.title}" 的稿件进行投稿前模拟评审。本轮由五个相互独立的评审角色完成：主编、方法学评审、领域专家评审、跨学科/实践视角评审和反方评审。

### 决策 / Decision: {rec_label(decision)}

### 评审报告概览 / Report Inventory

| 评审角色 / Reviewer | 建议 / Recommendation | 置信度 / Confidence | 主要问题 / Main Concerns |
|----------|----------------|------------|---------------|
{inventory_rows}

置信度说明：置信度是每个评审角色对自己判断可靠程度的自评，1 表示高度不确定，5 表示该角色认为证据充分且判断高度可靠。它不是论文通过概率，也不是模型准确率。

### 共识分析 / Consensus Analysis

- 评审小组的共同判断是：当前最优先的修改不是语言润色，而是把论文贡献、证据链和结论边界讲清楚。
- 方法学评审和反方评审提出的问题应优先处理，因为它们直接影响结论是否可被评审人接受。
- 领域评审与跨学科视角评审指向同一类修改：强化文献定位，同时在讨论影响时避免过度概括。

### 反方评审门槛 / Devil's Advocate Gate

{"反方评审检测到致命问题，因此最终决策不能为接收 / Accept。" if has_critical_da else "本地 MVP 规则未检测到致命反方问题，但主要反驳风险仍需要作者正面处理。"}

## 第二部分：修改路线图 / Part 2: Revision Roadmap

### 优先级 1：结构性修改 / Priority 1 — Must Fix
{p1}

### 优先级 2：内容补充 / Priority 2 — Should Fix
{p2}

### 优先级 3：文字与格式 / Priority 3 — Nice to Fix
{p3}

## 第三部分：建议修改顺序 / Part 3: Suggested Revision Sequence

1. 先重写贡献段，明确目标领域、最接近的竞争性研究和本文的新意。
2. 在调整结论前，先修复方法与证据链。
3. 补充局限性、替代解释和利益相关者影响。
4. 等论证稳定后，再统一题目、摘要、结论、参考文献和格式。

## 实现说明 / Implementation Note

本综合意见由本地确定性 MVP 规则引擎生成。若在后端连接 OpenAI-compatible 模型，系统会读取现有 agent prompt 文件，生成更深入的自然语言评审。
"""
    return {
        "decision": decision,
        "confidence": confidence,
        "must_fix_count": len(priority_1),
        "should_fix_count": len(priority_2),
        "markdown": markdown,
    }


def profile_to_dict(profile: ManuscriptProfile) -> dict[str, Any]:
    return {
        "title": profile.title,
        "abstract": profile.abstract,
        "word_count": profile.word_count,
        "reference_count": profile.reference_count,
        "section_names": profile.section_names,
        "language": profile.language,
        "primary_discipline": profile.primary_discipline,
        "secondary_disciplines": profile.secondary_disciplines,
        "research_paradigm": profile.research_paradigm,
        "methodology_type": profile.methodology_type,
        "target_tier": profile.target_tier,
        "maturity": profile.maturity,
        "signals": profile.signals,
    }


def run_llm_review(
    text: str,
    profile: ManuscriptProfile,
    mode: str,
    model: str,
    base_url: str,
    api_key: str,
    disable_proxy: bool = False,
    verify_ssl: bool = True,
    ca_bundle: str = "",
    progress: Any = None,
) -> dict[str, Any]:
    # OpenAI-compatible, dependency-free HTTP client. Supports OPENAI_BASE_URL for
    # local gateways or other compatible providers.
    base_url = (base_url or "https://api.openai.com/v1").rstrip("/")
    model = model or "gpt-4o-mini"
    field_prompt = load_agent_prompt("field_analyst")
    emit_progress(progress, "field_analyst", "running", message="正在分析领域和配置评审团队")
    field_report = call_chat_completion(
        base_url=base_url,
        api_key=api_key,
        disable_proxy=disable_proxy,
        verify_ssl=verify_ssl,
        ca_bundle=ca_bundle,
        model=model,
        messages=[
            {"role": "system", "content": field_prompt},
            {
                "role": "user",
                "content": (
                    "Output language requirement: write all concrete review comments, critique, questions, and revision suggestions in Chinese. "
                    "You may keep headings, manuscript title, role names, and technical terms in Chinese-English bilingual form.\n\n"
                    f"Review mode: {mode}\n\nManuscript:\n{text[:60000]}"
                ),
            },
        ],
    )
    emit_progress(progress, "field_analyst", "complete", recommendation=rec_label("Panel configured"), confidence=4)

    reviewer_results = []
    for role in REVIEWER_ROLES:
        prompt = load_agent_prompt(role)
        emit_progress(progress, role, "running", message=f"{ROLE_LABELS[role]} 正在调用 LLM")
        content = call_chat_completion(
            base_url=base_url,
            api_key=api_key,
            disable_proxy=disable_proxy,
            verify_ssl=verify_ssl,
            ca_bundle=ca_bundle,
            model=model,
            messages=[
                {"role": "system", "content": prompt},
                {
                    "role": "user",
                    "content": (
                        "Output language requirement: write all concrete review comments, critique, questions, and revision suggestions in Chinese. "
                        "You may keep headings, manuscript title, role names, and technical terms in Chinese-English bilingual form.\n\n"
                        f"Reviewer configuration:\n{field_report}\n\n"
                        f"Produce only your assigned review report. Manuscript:\n{text[:60000]}"
                    ),
                },
            ],
        )
        recommendation = extract_recommendation(content)
        confidence = extract_confidence(content)
        reviewer_results.append(
            {
                "id": role,
                "label": ROLE_LABELS[role],
                "status": "complete",
                "recommendation": recommendation,
                "confidence": confidence,
                "markdown": content,
            }
        )
        emit_progress(progress, role, "complete", recommendation=rec_label(recommendation), confidence=confidence)

    synth_prompt = load_agent_prompt("synthesizer")
    reviews_bundle = "\n\n---\n\n".join(f"# {r['label']}\n{r['markdown']}" for r in reviewer_results)
    emit_progress(progress, "synthesizer", "running", message="正在综合多位评审意见")
    synthesis_md = call_chat_completion(
        base_url=base_url,
        api_key=api_key,
        disable_proxy=disable_proxy,
        verify_ssl=verify_ssl,
        ca_bundle=ca_bundle,
        model=model,
        messages=[
            {"role": "system", "content": synth_prompt},
            {
                "role": "user",
                "content": (
                    "Output language requirement: write all concrete editorial synthesis, consensus analysis, and revision roadmap items in Chinese. "
                    "You may keep headings, manuscript title, role names, and technical terms in Chinese-English bilingual form.\n\n"
                    f"Field analysis:\n{field_report}\n\nReviewer reports:\n{reviews_bundle}\n\n"
                    "Synthesize the reports into the editorial decision package."
                ),
            },
        ],
    )
    decision = extract_recommendation(synthesis_md)
    synthesis_confidence = extract_confidence(synthesis_md)
    emit_progress(progress, "synthesizer", "complete", recommendation=rec_label(decision), confidence=synthesis_confidence)
    return {
        "provider": "openai",
        "mode": mode,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "profile": profile_to_dict(profile),
        "agents": [
            {
                "id": "field_analyst",
                "label": ROLE_LABELS["field_analyst"],
                "status": "complete",
                "recommendation": rec_label("Panel configured"),
                "confidence": 4,
                "markdown": field_report,
            },
            *reviewer_results,
            {
                "id": "synthesizer",
                "label": ROLE_LABELS["synthesizer"],
                "status": "complete",
                "recommendation": rec_label(decision),
                "confidence": synthesis_confidence,
                "markdown": synthesis_md,
            },
        ],
        "summary": {
            "decision": decision,
            "confidence": synthesis_confidence,
            "markdown": synthesis_md,
        },
    }


def call_chat_completion(
    base_url: str,
    api_key: str,
    model: str,
    messages: list[dict[str, str]],
    disable_proxy: bool = False,
    verify_ssl: bool = True,
    ca_bundle: str = "",
) -> str:
    base_url = normalize_base_url(base_url)
    body = json.dumps({"model": model, "messages": messages, "temperature": 0.2}).encode("utf-8")
    req = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=body,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    try:
        handlers = []
        if disable_proxy:
            handlers.append(urllib.request.ProxyHandler({}))
        context = build_ssl_context(verify_ssl=verify_ssl, ca_bundle=ca_bundle)
        if context is not None:
            handlers.append(urllib.request.HTTPSHandler(context=context))
        opener = urllib.request.build_opener(*handlers) if handlers else urllib.request
        with opener.open(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(explain_http_error(exc.code, detail, base_url, model)) from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(explain_url_error(str(exc.reason), base_url, disable_proxy)) from exc
    return data["choices"][0]["message"]["content"]


def explain_http_error(code: int, detail: str, base_url: str, model: str) -> str:
    if code in {401, 403}:
        return f"LLM API 鉴权失败（HTTP {code}）。请检查 API Key 是否正确、是否有权限访问模型 {model}。服务返回：{detail[:500]}"
    if code == 404:
        return f"LLM API 地址或模型不存在（HTTP 404）。请检查 Base URL 是否应为类似 {base_url}，以及模型名 {model} 是否正确。服务返回：{detail[:500]}"
    if code == 429:
        return f"LLM API 额度或频率受限（HTTP 429）。请检查余额、限流或稍后重试。服务返回：{detail[:500]}"
    if code in {500, 502, 503, 504}:
        return f"LLM API 服务端或网络网关暂时不可用（HTTP {code}）。如果你使用代理，可能是代理隧道不可用。服务返回：{detail[:500]}"
    return f"LLM request failed: HTTP {code}. {detail[:500]}"


def explain_url_error(reason: str, base_url: str, disable_proxy: bool) -> str:
    if "Tunnel connection failed: 503" in reason:
        proxy_hint = "你已勾选绕过系统代理，但仍然失败；请检查网络是否能直连该 API。" if disable_proxy else "建议勾选“绕过系统代理”再试，或检查 Clash/VPN/公司代理是否可用。"
        return (
            "连接外部 LLM API 失败：HTTPS 代理隧道返回 503 Service Unavailable。"
            f"这通常表示本机代理/VPN/网络网关无法连接到 {base_url}，不是论文评审程序本身的问题。{proxy_hint}"
        )
    if "timed out" in reason.lower():
        return f"连接外部 LLM API 超时。请检查 Base URL、网络、代理或模型服务状态：{base_url}"
    if "CERTIFICATE_VERIFY_FAILED" in reason:
        return (
            f"连接外部 LLM API 失败：SSL 证书校验失败。请求目标：{base_url}。"
            "这通常不是 API Key 问题，而是本机 Python 不信任该站点证书链、代理/VPN 替换了证书，或服务端证书链不完整。"
            "建议先尝试在 app/config.json 中设置 ca_bundle 为可信 CA 文件路径；仅本地调试时可临时设置 verify_ssl=false。"
        )
    return f"连接外部 LLM API 失败：{reason}。请检查 Base URL、网络代理和 API 服务状态：{base_url}"


def build_ssl_context(verify_ssl: bool, ca_bundle: str) -> ssl.SSLContext | None:
    if not verify_ssl:
        return ssl._create_unverified_context()
    if ca_bundle:
        return ssl.create_default_context(cafile=ca_bundle)
    return None


def normalize_base_url(base_url: str) -> str:
    normalized = (base_url or "https://api.openai.com/v1").rstrip("/")
    suffix = "/chat/completions"
    if normalized.endswith(suffix):
        normalized = normalized[: -len(suffix)]
    return normalized.rstrip("/")


def extract_recommendation(markdown: str) -> str:
    patterns = [
        ("Reject", [r"\bReject\b", "拒稿", "拒绝", "退稿"]),
        ("Major Revision", [r"\bMajor Revision\b", "大修", "重大修改"]),
        ("Minor Revision", [r"\bMinor Revision\b", "小修", "轻微修改"]),
        ("Accept", [r"\bAccept\b", "接收", "接受", "录用"]),
    ]
    for rec, rec_patterns in patterns:
        if any(re.search(pattern, markdown, re.I) for pattern in rec_patterns):
            return rec
    return "Review Complete"


def extract_confidence(markdown: str) -> int:
    match = re.search(r"(?:confidence(?: score)?|置信度)\D+([1-5])", markdown, re.I)
    return int(match.group(1)) if match else 4
